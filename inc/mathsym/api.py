from workfolder import get_model,  storage, login_required_auth
from flask import flash,Blueprint, current_app, redirect, render_template, request, \
    session, url_for,send_file,Flask,Response,send_from_directory
import os
import zipfile
import pymysql
import pymysql.cursors
from urllib.parse import quote
from docx import Document

from openpyxl.workbook import Workbook
from openpyxl import load_workbook   
from io import BytesIO

import config

crudapi = Blueprint('crudapi', __name__)


@crudapi.route("/stat/downloaddocx", methods=['GET', 'POST'])
def statDownloadDocx():
    # Create a new Document
    document = Document()
    
    # Add content to the document
    document.add_heading('My Document', 0)
    document.add_paragraph('This is a simple document generated in Flask.')
    
    # Add a table
    table = document.add_table(rows=3, cols=3)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"Row {i+1}, Col {j+1}"
    
    # Save document to a BytesIO buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    # Create response
    return Response(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-disposition": "attachment; filename=my_document.docx"}
    )


@crudapi.route("/stat/downloadxlsx", methods=['GET', 'POST'])
def statDownloadxlsx():
    """
    Export Xlsx File

    Parameters:

    data (json) : request.get_json() 

    Returns:

    stream : return file buffer stream.

    """ 
    if request.method == 'POST':
        #data = request.data.decode('utf-8')
        data = request.get_json()
        #print("Received data:", data['sheet'])        
        sheetnames = list(data.keys())
        wb = Workbook()
        ws = wb.active
        ws.title = sheetnames[0]
        for idx,sht in enumerate(sheetnames):
            print(idx,sht)
            if idx==0:
                pass
            else:
                ws = wb.create_sheet(sht)          
            for row in data[sht]:
                if str(type(row))=="<class 'list'>":
                    ws.append(row)        
                elif str(type(row))=="<class 'str'>":
                    ws.append([row])        
                elif str(type(row))=="<class 'dict'>":
                    ws.append([elm[1] for elm in  row.items()])        
                else:    
                    print(str(type(row)))
                    ws.append(list(row))        
            
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        # Create response
        return Response(
            buffer,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-disposition": "attachment; filename=my_data.xlsx"}
        )

    # Create a new workbook
    wb = Workbook()
    # Get the active worksheet
    ws = wb.active
    ws.title = "My Sheet"
    # Sample data
    data = [
        ["Name", "Age", "Occupation"],
        ["John Doe", 30, "Engineer"],
        ["Jane Smith", 25, "Designer"],
        ["Bob Johnson", 40, "Manager"]
    ]
    # Write data to worksheet
    for row in data:
        ws.append(row)
    # Style the header row
    for cell in ws[1]:
        cell.font = cell.font.copy(bold=True)
    # Save to BytesIO buffer
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    # Create response
    return Response(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-disposition": "attachment; filename=my_data.xlsx"}
    )